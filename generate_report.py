from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()

    # Title Page
    title = document.add_heading('E-Commerce Deep Learning Microservice Architecture', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('\n')
    subtitle = document.add_paragraph('Final Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_page_break()

    # 1. Executive Summary
    document.add_heading('1. Executive Summary', level=1)
    document.add_paragraph(
        "This project outlines the development and deployment of a production-ready, Deep Learning-based "
        "Customer Intelligence Platform designed for the retail and e-commerce sector. Transitioning away from "
        "legacy machine learning algorithms, the system leverages PyTorch Neural Networks to drive three primary "
        "business functions: predictive Customer Lifetime Value (CLV) forecasting, churn risk classification, "
        "and hyper-personalized product recommendation."
    )
    document.add_paragraph(
        "To guarantee high-performance real-time execution, the models are encapsulated within a containerized "
        "FastAPI microservice architecture. This backend seamlessly feeds a Streamlit dashboard, providing stakeholders "
        "with real-time insights, interactive 'What-If' simulation capabilities, and explainable AI heuristics."
    )

    # 2. System Architecture
    document.add_heading('2. System Architecture & Engineering', level=1)
    document.add_paragraph(
        "The system abandons monolithic design principles in favor of a scalable microservice infrastructure. "
        "The architecture is composed of two primary nodes:"
    )
    
    document.add_heading('2.1 Backend: FastAPI Inference Engine', level=2)
    document.add_paragraph(
        "The backend is a RESTful API built on FastAPI, serving PyTorch neural networks in real-time. "
        "It decouples model training from inference, ensuring O(1) latency lookup speeds and preventing server "
        "bottlenecks during high user traffic. Artifacts (.pt weights and .pkl scaling pipelines) are loaded into "
        "memory upon container instantiation."
    )
    
    document.add_heading('2.2 Frontend: Streamlit Command Center', level=2)
    document.add_paragraph(
        "The frontend is a dynamic Streamlit web application. It executes HTTP requests to the FastAPI backend to "
        "render deep learning predictions. It features dynamic customer cohort badging (e.g., 'At-Risk VIP', 'Champion'), "
        "interactive parameter manipulation sliders, and batch CSV processing capabilities."
    )

    # 3. Deep Learning Modeling
    document.add_heading('3. Deep Learning Methodology', level=1)
    
    document.add_heading('3.1 Multi-Task Learning (MTL) for CLV & Churn', level=2)
    document.add_paragraph(
        "Instead of maintaining separate models for churn and lifetime value, this project implements a Multi-Task "
        "Neural Network. A shared hidden layer extracts dense representations from standard Recency, Frequency, "
        "and Monetary (RFM) inputs. The network then splits into:"
    )
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Regression Head: ").bold = True
    p.add_run("Predicts continuous 90-Day Customer Lifetime Value using Mean Squared Error loss.")
    p2 = document.add_paragraph(style='List Bullet')
    p2.add_run("Classification Head: ").bold = True
    p2.add_run("Outputs a sigmoid-activated Churn Probability using Binary Cross-Entropy loss.")

    document.add_heading('3.2 Neural Collaborative Filtering (NCF)', level=2)
    document.add_paragraph(
        "To replace traditional Market Basket Analysis algorithms like Apriori, an NCF architecture was deployed. "
        "The PyTorch model learns high-dimensional vector embeddings for both Users and Items. A dot-product mapping "
        "operation is executed to calculate predicted user-item purchase affinities. The raw logits are Min-Max scaled "
        "to a 75%-99% confidence range for highly interpretable frontend rendering."
    )

    # 4. Business Value & Prescriptive Analytics
    document.add_heading('4. Business Value & Prescriptive Analytics', level=1)
    document.add_paragraph(
        "A critical component of this project is the translation of raw machine learning outputs into actionable "
        "business logic. The application features a Prescriptive CRM Engine that issues automated directives. For example:"
    )
    p3 = document.add_paragraph(style='List Bullet')
    p3.add_run("High Churn Risk + High CLV: ").bold = True
    p3.add_run("Triggers an automated high-urgency win-back campaign with a 15% discount on top affinity items.")
    p4 = document.add_paragraph(style='List Bullet')
    p4.add_run("Low Churn Risk + High CLV: ").bold = True
    p4.add_run("Classifies the user as a 'Loyal Champion' and restricts discounting to protect profit margins, opting instead for VIP early-access workflows.")
    
    document.add_paragraph(
        "Furthermore, an Explainable AI (XAI) expander module isolates the primary heuristic factors driving "
        "the neural network's decision, ensuring that marketing managers maintain interpretability and trust in the system."
    )

    # 5. Conclusion
    document.add_heading('5. Conclusion', level=1)
    document.add_paragraph(
        "This project successfully bridges the gap between theoretical deep learning and practical, industry-ready "
        "software engineering. By leveraging PyTorch, FastAPI, and Docker, the system is highly scalable, incredibly "
        "accurate, and immediately deployable to a cloud environment."
    )

    document.save('ECommerce_Deep_Learning_Project_Report.docx')

if __name__ == "__main__":
    create_report()
