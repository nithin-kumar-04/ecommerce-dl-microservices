from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title_page(doc):
    doc.add_paragraph('\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CHAITANYA ENGINEERING COLLEGE')
    run.bold = True
    run.font.size = Pt(20)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Chaitanya Valley, Kommadi, Madhurawada, Visakhapatnam, Andhra Pradesh Pincode: 530048\n')
    p.add_run('Approved by AICTE New Delhi & Affiliated By JNTU Gurajada Vizianagaram\n')
    p.add_run('NAAC Accredited')
    
    doc.add_paragraph('\n\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FINAL YEAR PROJECT REPORT\nON\nE-COMMERCE DEEP LEARNING MICROSERVICE ARCHITECTURE')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph('\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Submitted in partial fulfillment of the requirements for the award of the degree of\n')
    run = p.add_run('Bachelor of Technology\n')
    run.bold = True
    p.add_run('In\n')
    run = p.add_run('Artificial Intelligence and Data Science')
    run.bold = True
    
    doc.add_paragraph('\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Name: ').bold = True
    p.add_run('Gudabandi Nithin Kumar\n')
    p.add_run('Roll Number: ').bold = True
    p.add_run('23L61A5418')
    
    doc.add_paragraph('\n\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department Of Artificial Intelligence')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_page_break()

def add_declaration(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('STUDENT DECLARATION')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph('_' * 65)
    
    doc.add_paragraph(
        'I, Gudabandi Nithin Kumar, bearing Roll Number 23L61A5418, a student of the '
        'Department of Artificial Intelligence and Data Science at Chaitanya Engineering College '
        '(JNTU-GV), Visakhapatnam, hereby declare that the final year project report titled '
        '"E-Commerce Deep Learning Microservice Architecture" is a true and original record '
        'of the practical work completed entirely by me.'
    )
    doc.add_paragraph(
        'I further declare that the knowledge and practical skills documented in this report were '
        'gained through my direct participation in building the Deep Learning PyTorch models, '
        'FastAPI microservices, and Streamlit dashboards. This comprehensive report is being '
        'submitted to the Department of Artificial Intelligence and Data Science to fulfill the '
        'academic requirements necessary for the award of my Bachelor of Technology degree.'
    )
    
    doc.add_paragraph('\n\nDate: ________________________')
    doc.add_paragraph('Signature: ____________________________')
    p = doc.add_paragraph()
    p.add_run('Name: ').bold = True
    p.add_run('Gudabandi Nithin Kumar')
    
    doc.add_page_break()

def add_certificate(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CERTIFICATE OF BONAFIDE WORK')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph('_' * 65)
    
    doc.add_paragraph(
        'This document is to certify that Gudabandi Nithin Kumar, a student of B.Tech in '
        'Artificial Intelligence and Data Science at Chaitanya Engineering College (JNTU-GV), '
        'Visakhapatnam, bearing Roll Number 23L61A5418, has successfully and diligently '
        'completed the final year engineering project titled "E-Commerce Deep Learning Microservice Architecture".'
    )
    doc.add_paragraph(
        'This written report accurately represents the student\'s genuine software development process, '
        'technical skill application, and practical implementation of Machine Learning and DevOps '
        'tools. The student demonstrated commendable dedication to understanding complex '
        'artificial intelligence concepts and applying them to real-world business scenarios.'
    )
    
    doc.add_paragraph('\n\n\nSignature: ________________________\t\tSignature: ________________________')
    doc.add_paragraph('Name: Mrs. J. Kavitha\t\t\t\tName: Mrs. Dr. K.N.S. Lakshmi')
    doc.add_paragraph('(Internal Guide)\t\t\t\t\t(Head of Department)')
    
    doc.add_page_break()

def add_abstract(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ABSTRACT')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph('_' * 65)
    
    doc.add_paragraph(
        "In today's highly competitive and data-driven corporate environment, businesses generate "
        "massive volumes of customer data. This comprehensive project report documents the development "
        "and deployment of a production-ready, Deep Learning-based Customer Intelligence Platform "
        "designed for the retail and e-commerce sector."
    )
    doc.add_paragraph(
        "Transitioning away from legacy machine learning algorithms, the system leverages PyTorch Neural "
        "Networks to drive three primary business functions: predictive Customer Lifetime Value (CLV) "
        "forecasting, churn risk classification, and hyper-personalized product recommendation using "
        "Neural Collaborative Filtering (NCF) and Multi-Task Learning (MTL)."
    )
    doc.add_paragraph(
        "To guarantee high-performance real-time execution, the models are encapsulated within a "
        "containerized FastAPI microservice architecture. This backend seamlessly feeds a Streamlit "
        "dashboard, providing stakeholders with real-time insights, interactive 'What-If' simulation "
        "capabilities, and explainable AI heuristics."
    )
    doc.add_page_break()

def add_content(doc):
    doc.add_heading('CHAPTER 1: INTRODUCTION', level=1)
    doc.add_paragraph(
        "The primary objective of this project is to achieve practical mastery over deep learning "
        "and production engineering. The curriculum was designed to confidently step into a "
        "corporate environment and immediately begin contributing to AI-driven projects."
    )
    
    doc.add_heading('CHAPTER 2: SYSTEM ARCHITECTURE', level=1)
    doc.add_paragraph(
        "The system abandons monolithic design principles in favor of a scalable microservice infrastructure. "
        "The backend is a RESTful API built on FastAPI, serving PyTorch neural networks in real-time. "
        "It decouples model training from inference. The frontend is a dynamic Streamlit web application."
    )
    
    doc.add_heading('CHAPTER 3: DEEP LEARNING METHODOLOGY', level=1)
    doc.add_paragraph(
        "Instead of maintaining separate models for churn and lifetime value, this project implements a Multi-Task "
        "Neural Network. A shared hidden layer extracts dense representations from standard Recency, Frequency, "
        "and Monetary (RFM) inputs. For product recommendations, a PyTorch Neural Collaborative Filtering (NCF) "
        "architecture was deployed to learn high-dimensional vector embeddings for Users and Items."
    )
    
    doc.add_heading('CHAPTER 4: CONCLUSION AND FUTURE SCOPE', level=1)
    doc.add_paragraph(
        "This project successfully bridges the gap between theoretical computer science concepts "
        "studied at Chaitanya Engineering College and the rigorous, practical software workflows "
        "demanded by modern corporate employers. Future scope involves integrating advanced "
        "statistical programming using Python frameworks to build intelligent autonomous systems."
    )

def main():
    doc = Document()
    add_title_page(doc)
    add_declaration(doc)
    add_certificate(doc)
    add_abstract(doc)
    add_content(doc)
    doc.save('Final_Year_Project_Report_Formatted.docx')

if __name__ == "__main__":
    main()
